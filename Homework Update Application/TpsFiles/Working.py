import urllib.request
import re
from selenium import webdriver
from pprint import pprint
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from bs4 import BeautifulSoup
import datetime
import sys
import time
import requests
import os
from Fromfolder import *
import xlwt
from docx import Document
from docx.shared import Inches,Pt

username = "enter here"
password = "enter here"
link = "https://topperseducation.sg/index.php"


DEBUG = False
optratio = 0.5

def dprint(text):
    if DEBUG is True:
        pprint(text)

#logging in
def login():
    #incognito part
    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_argument("--incognito")
    driver = webdriver.Chrome(ChromeDriverManager().install(),chrome_options = chrome_options)
    #next part
    driver.get(link)
    driver.find_elements_by_class_name("cancel")[0].click()
    driver.find_element_by_name("username").send_keys(username)
    driver.find_element_by_name("password").send_keys(password)
    driver.find_element_by_name("do-login").click()
    return driver


#hovering over
def hoveroverws(driver):
    element_to_hover_over = driver.find_element_by_xpath("/html/body/div/nav/div[2]/div/ul/li[4]/a/span")
    hover = ActionChains(driver).move_to_element(element_to_hover_over)
    hover.perform()


#collecting all data from upload ws
def collectfromclass(driver,key):
    html_content = driver.page_source
    soup = BeautifulSoup(html_content,'lxml')
    toptable = soup.find("table", attrs={"id": "order-listing"})
    topdata = toptable.tbody.find_all("tr")  
    # Get all the headings of Lists
    alldict = {}
    tr = 1
    for elem in topdata:
        i = 0
        s = str(elem.button)
        #print(s)
        x = s.lstrip('<button class="btn btn-outline-primary btn-upload-file" data-target="#modal-file-upload" data-toggle="modal" data-ws-id=')
        FinalID = x[:5]
        #print(f'data id onwards: {FinalID,type(FinalID)}')
        #print(FinalID)
        try:
            int(FinalID)
        except:
            FinalID = 'Not Submitted'
        d1 = {0:FinalID}
        for td in elem:
            #print(i,td.text,'\n')
            if i==9:
                #THIS IS CORRECTED/CORRECTION PENDING
                for k in td:
                    #print(k.text)
                    d1['corstat'] = k.text
                break
            if i==7:
                for k in td:
                    #print(k.text)
                    d1['wsstat'] = k.text
                    break
            d1[4] = tr
            if i==3:
                val = td.get("style")
                #print(f'style is {val}')
                if val == 'color:red':
                    #print('\nReached here\n')
                    d1[0] = 'Absent Last Class'
                #print(f'style is {td.get("style")}')
                d1[2] = td.text
            i+=1
        tr+=1
        #pprint(d1)
        if alldict.get(key,-1)==-1:
            alldict[key] = [(d1[2],d1[0],d1[4],d1['corstat'],d1['wsstat'])]
        else:
            alldict[key].append((d1[2],d1[0],d1[4],d1['corstat'],d1['wsstat']))
    return alldict


def collectfromup(driver):
##    driver.implicitly_wait(20)
##    link = driver.find_element_by_link_text('Upload Corrected Worksheet')
##    link.click()
##    select = Select(driver.find_element_by_name('order-listing_length'))
##    select.select_by_value('-1')
    html_content = driver.page_source

    soup = BeautifulSoup(html_content,'lxml')

    toptable = soup.find("table", attrs={"id": "order-listing"})
    topdata = toptable.tbody.find_all("tr")  

    # Get all the headings of Lists
    alldict = {}
    for elem in topdata:
        i = 0
        s = str(elem.button)
        l = s.split('id')
        m = l[1].lstrip('="')
        FinalID = m[:5]
        #print(FinalID)
        d1 = {0:FinalID}
        for td in elem:
            #print(i,td.text,'\n')
            if i==11:
                d1[1] = td.text
            if i==3:
                d1[2] = td.text
            if i==5:
                d1[3] = td.text
            i+=1
        if alldict.get(d1[1],-1)==-1:
            alldict[d1[1]] = [(d1[2],d1[3],d1[0])]
        else:
            alldict[d1[1]].append((d1[2],d1[3],d1[0]))
    return alldict

def collectfromview(driver):
    html_content = driver.page_source

    soup = BeautifulSoup(html_content,'lxml')

    toptable = soup.find("table", attrs={"id": "order-listing"})
    topdata = toptable.tbody.find_all("tr")  

    # Get all the headings of Lists
    alldict = {}
    for elem in topdata:
        #3 is batch, 5 is subject\nworksheet, 9 is date 
        i = 0

        s = str(elem.button)
        l = s.split('id')
        m = l[1].lstrip('="')
        FinalID = m[:5]
        d1 = {'FID':FinalID}

        for td in elem:
            #print(i,td.text,'\n')
            if i==3:
                txt = td.text
                txt = txt.lstrip('\n').split('\n')[0]
                d1['Batch'] = txt
            if i==5:
                txt = td.text
                txt = txt.lstrip('\n').split('\n')[0]
                d1['Subject'] = txt
            if i==9:
                d1['Date'] = td.text
            i+=1
        if alldict.get(d1['Date'],-1)==-1:
            alldict[d1['Date']] = [(d1['Batch'],d1['Subject'],d1['FID'])]
        else:
            alldict[d1['Date']].append((d1['Batch'],d1['Subject'],d1['FID']))
    return alldict



def stillpending(dates):
    driver = login()
    hoveroverws(driver)
    driver.implicitly_wait(20)
    link = driver.find_element_by_link_text('Upload Corrected Worksheet')
    link.click()
    select = Select(driver.find_element_by_name('order-listing_length'))
    select.select_by_value('-1')
    dtc = collectfromup(driver)
    driver.quit()
    flist = list(filter(lambda x:x[0] in dates,dtc))
    return flist
    
def gettingdict(dates):
    if stillpending(dates):
        print('Uh oh... Correction not yet complete...')
        time.sleep(2)
        return False
    driver = login()
    hoveroverws(driver)
    driver.implicitly_wait(20)
    button = WebDriverWait(driver, 15).until(EC.element_to_be_clickable((By.LINK_TEXT,'View Corrected Worksheet'))).click()
    select = Select(driver.find_element_by_name('order-listing_length'))
    select.select_by_value('-1')
    dtc2 = collectfromview(driver)
    #dprint(f'dtc2 here is {dtc2}')
    
    datetoclass2 = {}
    for (key,value) in dtc2.items():
        if key in dates:
            datetoclass2[key] = value
    dprint(f'date to class dictionary {datetoclass2}')
    
    #NOTE: HERE DRIVER IS AT VIEW WORKSHEET INDEX
            
    return (datetoclass2,driver)

def sort(dates):
    mdict = {'January':1,'February':40,'March':80,'April':120,'May':160,'June':200,'July':240,
             'August':280,'September':320,'October':360,'November':400,'December':440}
    sdates = []
    for date in dates:
        parts = date.split(' ')
        amt = mdict[parts[0]] + int(parts[1][:-1]) + int(parts[2])
        sdates.append((date,amt))
    sdates = sorted(sdates,key = lambda x:x[1])
    flist = list(map(lambda x:x[0],sdates))
    return flist


def analysis(datetoclass2,driver):
    def collectratio(driver):
        html_content = driver.page_source

        soup = BeautifulSoup(html_content,'lxml')

        toptable = soup.find("table", attrs={"id": "order-listing"})
        topdata = toptable.tbody.find_all("tr")  

        # Get all the headings of Lists
        students = []
        for elem in topdata:
            #3 is name, 7 is whether ws yes or no, 9 is corrected(not needed)
            #11 is comp/alc/incomp etc. 13 is \ncomments
            i = 0
            
            d1 = {}

            for td in elem:
                #print(i,td.text,'\n')
                if i==3:
                    d1['Name'] = td.text
                if i==7:
                    d1['WS Status'] = td.text
                if i==11:
                    d1['Correction'] = td.text
                if i==13:
                    txt = td.text
                    txt = txt.lstrip('\n').rstrip('\nView\nEdit')
                    d1['Comments'] = txt
                i+=1
            students.append(tuple(d1.values()))
        #dprint(f'students here are {students}')

        denom,num = len(students),0
        ignore = True
        
        for student in students:
            stat,corr = student[1],student[2]
            comments = student[3]
            #if completed etc. dont ignore this class
            if corr:
                ignore = False
            if stat == 'Worksheet Uploaded ':
                    if corr!='Not Submitted':
                        num+=1
            elif corr=='Completed' and comments!='LATE SUBMISSION':
                num+=1
        ratio = num/denom
        #dprint(f'ratio here is {ratio}')
        if ratio >= optratio or ignore:
            return -1
        elif ratio==0.0:
            return ''
        return f" only {num}/{denom}"
    
    reportdict = {}
    p = driver.current_window_handle
    for (date,classes) in datetoclass2.items():
        reportdict[date] = ([],[])
        for clas in classes:
            wsmid = clas[2]
            button = driver.find_element_by_xpath(f"//*[@data-wsm-id={wsmid}]").click()
            #chwd is child window
            chwd = driver.window_handles
            for w in chwd:
                if(w!=p):
                    driver.switch_to.window(w)
                    break
            #inside the class now
            select = Select(driver.find_element_by_name('order-listing_length'))
            select.select_by_value('-1')
            ratio = collectratio(driver)
            finalstring = f'{clas[0]}  {clas[1]}\t\t{ratio}\n'
            #print(f'final ratio up here is {finalstring}')
            if ratio!=-1:
                if ratio == '':
                    reportdict[date][0].append(finalstring)
                else:
                    reportdict[date][1].append(finalstring)
            driver.switch_to.window(p)
    #dprint(f'final report dict {reportdict}')
    driver.quit()
    return reportdict

#def f1(pname):
    #return (dates,pname)

def createreport(reportdict,dates,pname):
    mvaldict = {'January':'01','February':'02','March':'03','April':'04','May':'05','June':'06','July':'07',
             'August':'08','September':'09','October':'10','November':'11','December':'12'}
    def moddate(date):
        parts = date.split(' ')
        datetime_object = datetime.datetime.strptime(mvaldict[parts[0]], "%m")
        smo = datetime_object.strftime("%b")
        return f'{parts[1][:-1]}-{smo}-{parts[2]}'
        
    #WEEKLY REPORT 04-12-21
    enddate = dates[-1]
    parts = enddate.split(' ')
    path = pname + r'\WEEKLY REPORT ' + parts[1][:-1] + '-' + mvaldict[parts[0]] + '-' + parts[2][2:] + '.docx'
    if os.path.exists(path):
        os.remove(path)
        dprint('EX file removed')
    document = Document()
    
    sec_pr = document.sections[0]._sectPr # get the section properties el
    # create new borders el
    pg_borders = OxmlElement('w:pgBorders')
    # specifies how the relative positioning of the borders should be calculated
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single') # a single line
        border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), 'auto')
        pg_borders.append(border_el) # register single border to border el
    sec_pr.append(pg_borders)
    #document.add_heading('Document Title', 0)

    obj_styles = document.styles
    
    obj_mainstyle = obj_styles.add_style('Main', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_mainstyle.font
    obj_font.size = Pt(20)
    obj_font.bold = True
    obj_font.name = 'Browallia New'

    obj_classstyle = obj_styles.add_style('Classes', WD_STYLE_TYPE.CHARACTER)
    obj_font2 = obj_classstyle.font
    obj_font2.size = Pt(16)
    obj_font2.name = 'Cambria'
    obj_font2.bold = False

    opening = document.add_paragraph()
    #opening.style = document.styles['Main']
    opening.add_run(f'WEEKLY REPORT – {moddate(dates[0])} to {moddate(dates[-1])}',style = 'Main')
    opening.add_run('\n\n* LESS STUDENT SUBMISSIONS – Only few students submitted\n* ZERO STUDENT SUBMISSIONS – No student submitted',style = 'Main')
    opening.add_run('\n\nNOTE: REPORT IS BASED ON THE HW SUBMISSIONS ',style = 'Main')
    opening.add_run('IN PERSON',style = 'Main').underline = True
    opening.add_run(', AND ',style = 'Main')
    opening.add_run('ON TIME',style = 'Main').underline = True
    opening.add_run(' SPECIFICALLY \n--\n',style = 'Main')

    for date in dates:
        report = reportdict[date]
        zerosubs,lessubs = report[0],report[1]
        para = document.add_paragraph()
        para.add_run(f'{moddate(date)}',style = 'Main').underline = True
        para.add_run('\nZERO STUDENT SUBMISSIONS',style = 'Main')
        cls = document.add_paragraph()
        if zerosubs:
            zerosubs[-1] = zerosubs[-1].rstrip('\n')
            for x in zerosubs:
                cls.add_run(x,style='Classes')
        else:
            cls.add_run('-',style='Classes')
        para2 = document.add_paragraph()
        para2.add_run('LESS STUDENT SUBMISSIONS',style = 'Main')
        cls2 = document.add_paragraph()
        if lessubs:
            lessubs[-1] = lessubs[-1].rstrip('\n')
            for x in lessubs:
                cls2.add_run(x,style='Classes')
            #para.add_run(f'{[x for x in lessubs]}',style = 'Classes')
        else:
            cls2.add_run('-',style='Classes')
        para3 = document.add_paragraph()
        para3.add_run('\n')
        
    
    document.save(path)

def corefn(pname):
    dates = getdates(pname)
    dates = sort(dates)
    tup = gettingdict(dates)
    if tup is False:
        return
    datetoclass2,driver = tup[0],tup[1]
    reportdict = analysis(datetoclass2,driver)
    createreport(reportdict,dates,pname)
    

#corefn()
